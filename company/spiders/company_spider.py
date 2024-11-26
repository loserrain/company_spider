# spiders/company_spider.py
import scrapy
import json
import re
import os
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class CompanySpider(scrapy.Spider):
    name = 'company_spider'
    allowed_domains = ['104.com.tw']
    
    def __init__(self, *args, **kwargs):
        super(CompanySpider, self).__init__(*args, **kwargs)
        # 建立 output 資料夾
        self.output_dir = 'output'
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
    
    def start_requests(self):
        company_id = 'e6o7g3l'
        api_url = f'https://www.104.com.tw/company/ajax/content/{company_id}?resourceStatus=1'
        
        headers = {
            'Accept': 'application/json, text/plain, */*',
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Referer': f'https://www.104.com.tw/company/{company_id}',
        }
        
        yield scrapy.Request(
            url=api_url,
            headers=headers,
            callback=self.parse
        )

    def get_safe_filename(self, filename):
        """將公司名稱轉換為安全的檔名"""
        unsafe_chars = r'[<>:"/\\|?*]'
        filename = re.sub(unsafe_chars, '', filename)
        filename = filename.strip()
        return filename if filename else 'company'

    def parse(self, response):
        data = json.loads(response.text)
        company_data = data.get('data', {})
        
        company_name = company_data.get('custName', '')
        safe_filename = self.get_safe_filename(company_name)
        
        result = {
            '公司名稱': company_name,
            '統一編號': company_data.get('custNo', ''),
            '產業類別': company_data.get('industryDesc', ''),
            '產業分類': company_data.get('indcat', ''),
            '員工人數': company_data.get('empNo', ''),
            '資本額': company_data.get('capital', ''),
            '公司地址': company_data.get('address', ''),
            '公司網站': company_data.get('custLink', ''),
            '公司簡介': company_data.get('profile', ''),
            '主要商品': company_data.get('product', ''),
            '福利制度': company_data.get('welfare', ''),
            '經營理念': company_data.get('management', ''),
            '公司電話': company_data.get('phone', ''),
            '傳真號碼': company_data.get('fax', ''),
            '聯絡人': company_data.get('hrName', ''),
            '公司標籤': company_data.get('tagNames', []),
            '法定標籤': company_data.get('legalTagNames', []),
            '最新消息': {
                '標題': company_data.get('news', ''),
                '連結': company_data.get('newsLink', '')
            },
            '公司發展': company_data.get('historys', [])
        }
        
        self.generate_files(result, safe_filename)
        
        yield result

    def generate_files(self, data, filename):
        """統一處理所有檔案的生成"""
        filepath = os.path.join(self.output_dir, filename)
        
        # 生成JSON
        with open(f'{filepath}.json', 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        
        # 生成Markdown
        self.generate_markdown(data, filepath)
        
        # 生成Word
        self.generate_word(data, filepath)
        
        # 生成CSV
        self.generate_csv(data, filepath)

    def generate_csv(self, data, filepath):
        """生成CSV檔案"""
        # 準備基本資料
        basic_data = {
            '公司名稱': [data['公司名稱']],
            '統一編號': [data['統一編號']],
            '產業類別': [data['產業類別']],
            '產業分類': [data['產業分類']],
            '員工人數': [data['員工人數']],
            '資本額': [data['資本額']],
            '公司地址': [data['公司地址']],
            '公司網站': [data['公司網站']],
            '公司電話': [data['公司電話']],
            '傳真號碼': [data['傳真號碼']],
            '聯絡人': [data['聯絡人']],
            '公司標籤': [', '.join(data['公司標籤'])],
            '法定標籤': [', '.join(data['法定標籤'])],
            '最新消息標題': [data['最新消息']['標題']],
            '最新消息連結': [data['最新消息']['連結']]
        }
        
        # 建立DataFrame並保存
        df = pd.DataFrame(basic_data)
        df.to_csv(f'{filepath}.csv', index=False, encoding='utf-8-sig')

    def generate_markdown(self, data, filepath):
        markdown_content = f"""# {data['公司名稱']} 公司資料

## 基本資訊
- 統一編號：{data['統一編號']}
- 產業類別：{data['產業類別']}
- 產業分類：{data['產業分類']}
- 員工人數：{data['員工人數']}
- 資本額：{data['資本額']}
- 公司地址：{data['公司地址']}
- 公司網站：{data['公司網站']}
- 公司電話：{data['公司電話']}
- 傳真號碼：{data['傳真號碼']}
- 聯絡人：{data['聯絡人']}

## 公司簡介
{data['公司簡介']}

## 主要商品
{data['主要商品']}

## 福利制度
{data['福利制度']}

## 經營理念
{data['經營理念']}

## 公司標籤
{', '.join(data['公司標籤'])}

## 法定標籤
{', '.join(data['法定標籤'])}

## 最新消息
- 標題：{data['最新消息']['標題']}
- 連結：{data['最新消息']['連結']}

## 公司發展歷程
"""
        for history in data['公司發展']:
            markdown_content += f"### {history['year']}年{history['month']}月\n{history['content']}\n\n"

        with open(f'{filepath}.md', 'w', encoding='utf-8') as f:
            f.write(markdown_content)

    def generate_word(self, data, filepath):
        doc = Document()
        
        title = doc.add_heading(f"{data['公司名稱']} 公司資料", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading('基本資訊', 1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        basic_info = [
            ('統一編號', data['統一編號']),
            ('產業類別', data['產業類別']),
            ('產業分類', data['產業分類']),
            ('員工人數', data['員工人數']),
            ('資本額', data['資本額']),
            ('公司地址', data['公司地址']),
            ('公司網站', data['公司網站']),
            ('公司電話', data['公司電話']),
            ('傳真號碼', data['傳真號碼']),
            ('聯絡人', data['聯絡人'])
        ]
        
        for label, value in basic_info:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = str(value)
        
        sections = [
            ('公司簡介', data['公司簡介']),
            ('主要商品', data['主要商品']),
            ('福利制度', data['福利制度']),
            ('經營理念', data['經營理念'])
        ]
        
        for title, content in sections:
            doc.add_heading(title, 1)
            doc.add_paragraph(content)
        
        doc.add_heading('公司標籤', 1)
        doc.add_paragraph(', '.join(data['公司標籤']))
        
        doc.add_heading('法定標籤', 1)
        doc.add_paragraph(', '.join(data['法定標籤']))
        
        doc.add_heading('最新消息', 1)
        doc.add_paragraph(f"標題：{data['最新消息']['標題']}")
        doc.add_paragraph(f"連結：{data['最新消息']['連結']}")
        
        doc.add_heading('公司發展歷程', 1)
        for history in data['公司發展']:
            doc.add_heading(f"{history['year']}年{history['month']}月", 2)
            doc.add_paragraph(history['content'])
        
        doc.save(f'{filepath}.docx')

# settings.py
BOT_NAME = 'company'
SPIDER_MODULES = ['company.spiders']
NEWSPIDER_MODULE = 'company.spiders'

ROBOTSTXT_OBEY = True
DOWNLOAD_DELAY = 2

DEFAULT_REQUEST_HEADERS = {
   'Accept': 'application/json, text/plain, */*',
   'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
}